from docx import Document
from pathlib import Path
import json, re, shutil

root = Path(r'E:\桌面\考研英语定制化')
src = Path(r'C:\Users\30700\Downloads\考研英语作文模拟题30篇（含图片） (1)\kaoyan_essay')
doc = Document(src / '考研英语作文模拟题30篇.docx')
paras = [p.text.strip() for p in doc.paragraphs]
items = []
for i, text in enumerate(paras):
    m = re.match(r'第(\d+)题\s+【(.+?)】', text)
    if not m: continue
    n, label = int(m.group(1)), m.group(2)
    prompt = next((x for x in paras[i+1:i+5] if x.startswith('Directions:')), '')
    exam = '英语一' if n <= 15 else '英语二'
    part = '小作文' if (exam == '英语一' and n <= 8) or (exam == '英语二' and n <= 23) else '大作文'
    typ = '图画作文' if '图画' in label else ('图表作文' if ('图表' in label or '柱状图' in label or '折线图' in label or '饼图' in label) else ('通知' if '通知' in label else '书信'))
    image = None
    if 9 <= n <= 14: image = f'assets/simulations/pic{n-8:02d}.png'
    if n == 15: image = 'assets/simulations/chart08.png'
    if 24 <= n <= 30: image = f'assets/simulations/chart{n-23:02d}.png'
    q = {'id': f'source-sim-{n:02d}', 'exam': exam, 'part': part, 'type': typ, 'topic': label.split('·')[-1].rstrip('】'), 'title': f'资料模拟题 {n} · {label}', 'text': prompt, 'source': next((x for x in paras[i+1:i+7] if x.startswith('（来源：')), '本地资料包'), 'image': {'file': image} if image else None}
    items.append(q)

assets = root / 'assets' / 'simulations'; assets.mkdir(parents=True, exist_ok=True)
for p in (src / 'images').glob('*.png'):
    if p.name.startswith('pic'):
        dest = assets / ('pic' + p.name[3:5] + '.png')
    else:
        m = re.match(r'(\d+)_', p.name); dest = assets / ('chart' + f'{int(m.group(1)):02d}' + '.png') if m else assets / p.name
    shutil.copy2(p, dest)

(root / 'js' / 'data-simulations-source.js').write_text('window.APP_DATA_SIMULATIONS_SOURCE = ' + json.dumps(items, ensure_ascii=False, indent=2) + ';\n', encoding='utf-8')

phrase_doc = Document(Path(r'C:\Users\30700\Downloads\考研英语作文好词好句汇总 (1)\考研英语作文好词好句汇总.docx'))
phrases = []
for text in [p.text.strip() for p in phrase_doc.paragraphs if p.text.strip()]:
    if re.search(r'[A-Za-z]{3}', text) and not text.startswith(('e.g.', 'Directions:')):
        parts = re.split(r'\s+→\s+|\s+—\s+|\s+：\s+', text, maxsplit=1)
        english = parts[0].strip()
        if len(english) > 3:
            phrases.append({'id': 'imported-' + str(len(phrases)+1), 'category': '资料汇总', 'text': english, 'translation': parts[1].strip() if len(parts) > 1 else '', 'usage': '来自本地好词好句汇总'})
(root / 'js' / 'data-phrases-imported.js').write_text('window.APP_DATA_PHRASES_IMPORTED = ' + json.dumps(phrases, ensure_ascii=False, indent=2) + ';\n', encoding='utf-8')
print(f'simulations={len(items)} phrases={len(phrases)} assets={len(list(assets.glob("*.png")))}')
